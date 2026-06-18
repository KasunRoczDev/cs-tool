#!/usr/bin/env python3
import sys
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x38, 0x64)
BLUE = RGBColor(0x2E, 0x54, 0x96)
GREY = RGBColor(0x88, 0x88, 0x88)

doc = Document()
st = doc.styles['Normal']
st.font.name = 'Arial'
st.font.size = Pt(10.5)

for lvl, sz, col in (('Heading 1', 15, NAVY), ('Heading 2', 12, BLUE)):
    s = doc.styles[lvl]
    s.font.name = 'Arial'; s.font.size = Pt(sz); s.font.bold = True; s.font.color.rgb = col

def shade(cell, hexv):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexv)
    tcPr.append(sh)

def set_cell(cell, text, bold=False, color=None, size=9):
    cell.text = ''
    p = cell.paragraphs[0]; r = p.add_run(str(text))
    r.font.size = Pt(size); r.font.bold = bold; r.font.name = 'Arial'
    if color: r.font.color.rgb = color

def add_table(headers, rows, widths):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.autofit = False
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        set_cell(hdr[i], h, bold=True, color=RGBColor(0xFF,0xFF,0xFF), size=9)
        shade(hdr[i], '1F3864')
    for ri, row in enumerate(rows):
        cells = t.add_row().cells
        for i, v in enumerate(row):
            set_cell(cells[i], v, size=9)
            if ri % 2: shade(cells[i], 'F2F5FA')
    for r in t.rows:
        for i, c in enumerate(r.cells):
            c.width = Inches(widths[i])
    doc.add_paragraph()
    return t

def bullet(text):
    doc.add_paragraph(text, style='List Bullet')

# Title page
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for _ in range(4): p.add_run('\n')
r = p.add_run('Monitoring & Security Reference'); r.bold=True; r.font.size=Pt(26); r.font.color.rgb=NAVY
p2 = doc.add_paragraph(); p2.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p2.add_run('DevOps / DevSecOps Metrics, Risks & Alert Thresholds'); r.font.size=Pt(14); r.font.color.rgb=RGBColor(0x55,0x55,0x55)
p3=doc.add_paragraph(); p3.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p3.add_run('Cybersecurity & Server Metrics Monitoring Platform  ·  June 2026'); r.font.size=Pt(10); r.font.color.rgb=GREY

doc.add_page_break()

doc.add_heading('1. How to use this reference', level=1)
doc.add_paragraph("This is the complete checklist of what to monitor across three layers — server health, security, and application — with concrete alert thresholds. It pairs with the monitoring/ add-on folder in this repository, which implements the gaps the platform did not previously cover. Rows marked NEW are the additions; the rest your agent already collects.")
doc.add_paragraph("Severity convention: critical = page immediately; high = investigate within the hour; medium = review same day; low = informational / trend only.")

doc.add_heading('2. Core server parameters (system health)', level=1)
doc.add_heading('2.1 CPU & load', level=2)
add_table(["Metric","Why it matters / risk","Threshold","Status"],[
 ["CPU usage % (user/sys/iowait)","Exhaustion → DDoS, infinite loop, bad deploy","warn 80%, crit 90% 5m","Have"],
 ["Load average 1/5/15","Trend before crash; 5/15 show direction","see load_per_core","5/15 NEW"],
 ["Load per core (load1 ÷ cores)",">1 = demand exceeds CPUs","warn 1.5, crit 3.0","NEW"],
 ["iowait %","High iowait = disk is bottleneck, not CPU","warn 20%","via CPU"],
],[2.0,2.7,1.6,0.6])
doc.add_heading('2.2 Memory', level=2)
add_table(["Metric","Why it matters / risk","Threshold","Status"],[
 ["Used / available memory %","Leak (Node/PHP-FPM) → crashes","warn 85%, crit 95%","Have"],
 ["Swap usage %","Swap thrashing = severe latency before OOM","warn 25%, crit 50%","NEW"],
 ["OOM-killer events","Kernel killed a process — silent outage","any = high","via auditd"],
],[2.0,2.7,1.6,0.6])
doc.add_heading('2.3 Disk', level=2)
add_table(["Metric","Why it matters / risk","Threshold","Status"],[
 ["Disk usage %","Full disk → DB crash, log loss","warn 80%, crit 90%","Have"],
 ["Inode usage %","Out of inodes = can't create files","warn 80%, crit 90%","NEW"],
 ["Disk read/write B/s","Saturation → DB + API latency","baseline + trend","NEW"],
 ["I/O queue depth","Sustained queue = slow disk","warn 8, crit 20","NEW"],
],[2.0,2.7,1.6,0.6])
doc.add_heading('2.4 Network & connections', level=2)
add_table(["Metric","Why it matters / risk","Threshold","Status"],[
 ["In/out bytes per sec","DDoS, exfil, saturation","baseline + spike","Have"],
 ["TCP ESTABLISHED","Connection load","baseline","NEW"],
 ["TCP TIME_WAIT","Pileup exhausts ephemeral ports","warn 20k, crit 40k","NEW"],
 ["TCP SYN_RECV","Sustained high = SYN flood","warn 256, crit 1024","NEW"],
 ["conntrack table %","Table full → silently dropped conns","warn 70%, crit 85%","NEW"],
 ["File descriptors %","High-connection servers die here first","warn 70%, crit 85%","NEW"],
],[2.0,2.7,1.6,0.6])
doc.add_heading('2.5 Process & host integrity', level=2)
add_table(["Metric","Why it matters / risk","Threshold","Status"],[
 ["Total / running processes","Fork bomb, runaway spawning","baseline + spike","NEW"],
 ["Zombie processes","Defunct children pile up","warn 20, crit 100","NEW"],
 ["Uptime","Unexpected reboot = crash or attack","drop = investigate","NEW"],
 ["NTP / time drift","Breaks TLS, auth, log correlation","warn 500ms, crit 2s","NEW"],
],[2.0,2.7,1.6,0.6])

doc.add_page_break()
doc.add_heading('3. Security monitoring (DevSec)', level=1)
doc.add_heading('3.1 Authentication & access', level=2)
add_table(["Signal","Risk","Threshold / action","Status"],[
 ["SSH failed logins","Brute force, credential stuffing","warn 20/min, crit 100/min","Have"],
 ["Root login attempts","Direct privileged compromise","any = high; disable root","Have"],
 ["sudo / su usage & failures","Privilege escalation","failures = high","Have"],
 ["GeoIP / impossible travel","Stolen creds from new country","new country = review","Gap*"],
 ["PermitRootLogin / PasswordAuth","Weak SSH config","flag if enabled","NEW"],
],[2.0,2.4,1.9,0.6])
doc.add_heading('3.2 Network & firewall', level=2)
add_table(["Signal","Risk","Threshold / action","Status"],[
 ["UFW / iptables blocks","Lateral movement, scanning","spike = medium","Have"],
 ["Firewall inactive / permissive","No perimeter","critical","Have"],
 ["Open dangerous ports (DB/Redis)","Exposed data services","any exposed = high","Have"],
 ["Outbound traffic anomaly","Malware C2, exfiltration","Falco rule + baseline","NEW"],
 ["Network IDS (Suricata/Zeek)","Packet-level intrusion","deploy Suricata","Gap*"],
],[2.0,2.4,1.9,0.6])
doc.add_heading('3.3 OS, process & file integrity', level=2)
add_table(["Signal","Risk","Tool","Status"],[
 ["New root process / suspicious exec","Backdoor, cryptominer","auditd + Falco","Have/NEW"],
 ["Binaries dropped in /tmp /dev/shm","Dropper, miner","Falco rule","NEW"],
 ["Shell spawned in container","Post-exploitation","Falco rule","NEW"],
 ["Changes to passwd/shadow/sudoers","Priv-esc, backdoor account","auditd rules","NEW"],
 ["cron / systemd unit changes","Persistence","auditd rules","NEW"],
 ["app .env / SSH key reads","Credential theft","Falco + auditd","NEW"],
 ["Kernel module load","Rootkit","auditd","NEW"],
],[2.1,2.3,1.9,0.6])
doc.add_heading('3.4 Vulnerability & supply chain', level=2)
add_table(["Signal","Risk","Tool","Status"],[
 ["TLS certificate expiry","Outage — very common","cert-cve.js / TLS script","NEW"],
 ["Container image CVEs","Shipping known-vuln images","Trivy in CI","NEW"],
 ["Dependency CVEs (npm audit)","Vulnerable libraries","cert-cve.js / Trivy fs","NEW"],
 ["Pending OS security updates","Unpatched CVE exposure","cert-cve.js / unattended","NEW"],
 ["Secrets in commits","Leaked keys/.env","gitleaks in CI","NEW"],
 ["Host hardening drift","Misconfiguration","Lynis audit","NEW"],
],[2.1,2.3,1.9,0.6])
doc.add_heading('3.5 Logs to monitor & protect', level=2)
bullet("/var/log/auth.log, syslog, kern.log — auth, sudo, kernel events.")
bullet("nginx/apache access + error logs — exploit probes, scanners, 5xx.")
bullet("/var/log/audit/audit.log — auditd integrity events.")
bullet("Application logs (NestJS, OMS, courier, IoT) — errors and anomalies.")
bullet("Critical: ship logs OFF-HOST so an attacker who roots the box cannot wipe the evidence.")

doc.add_page_break()
doc.add_heading('4. Application-level metrics', level=1)
doc.add_heading('4.1 API layer', level=2)
add_table(["Metric","Risk","Threshold","Status"],[
 ["Request rate (RPS)","Bot attack, retry storm","baseline + spike","Partial"],
 ["Error rate 4xx / 5xx","Broken deploy, hidden DB issue","warn 1%, crit 5% 5xx","Partial"],
 ["Latency p95 / p99","Slow API = hidden bottleneck","p95 < 300ms (SLO)","Partial"],
 ["Distributed tracing","Can't see time across services","OpenTelemetry","NEW"],
 ["Synthetic / uptime probe","Confirm 'actually up' externally","blackbox_exporter","NEW"],
],[2.0,2.4,1.9,0.6])
doc.add_heading('4.2 Queue (BullMQ / Redis)', level=2)
add_table(["Metric","Risk","Threshold","Status"],[
 ["Waiting (backlog)","Backlog explosion, silent delay","warn 1000","NEW"],
 ["Active = 0 with backlog","Stuck worker","any = high","NEW"],
 ["Failed jobs / DLQ","Lost work after retries","rate > 1%","NEW"],
],[2.0,2.4,1.9,0.6])
doc.add_heading('4.3 PHP-FPM (per pool)', level=2)
add_table(["Metric","Risk","Threshold","Status"],[
 ["Pool utilisation (active/total)","All workers busy → requests wait","warn 90%, crit 100%","NEW"],
 ["max_children reached count","Requests queued/dropped","any > 0 = high","NEW"],
 ["Listen queue depth","Backlog = user-visible latency","warn 1, crit 50","NEW"],
 ["Slow requests","Slow code/DB path","any > 0 = medium","NEW"],
 ["Top-CPU worker + its request","Which request burns CPU","worker >= 80% = flag","NEW"],
 ["Top-memory worker + its request","Which request burns memory","track + trend","NEW"],
],[2.0,2.4,1.9,0.6])
doc.add_heading('4.4 Database (PostgreSQL / TimescaleDB)', level=2)
add_table(["Metric","Risk","Threshold","Status"],[
 ["Connection usage %","Connection exhaustion","warn 75%, crit 90%","NEW"],
 ["Idle-in-transaction","Leaked txns, lock holds","trend up = investigate","NEW"],
 ["Blocked locks / deadlocks","Contention, stalls","any = review","NEW"],
 ["Long-running queries (>5s)","Query bottleneck","any = review","NEW"],
 ["Replication lag","Stale reads, failover risk","warn 30s, crit 120s","NEW"],
],[2.0,2.4,1.9,0.6])
doc.add_heading('4.5 Cache (Redis)', level=2)
add_table(["Metric","Risk","Threshold","Status"],[
 ["Hit / miss ratio","Cache cold → DB load","warn < 0.8, crit < 0.5","NEW"],
 ["Used memory","Memory pressure","vs maxmemory","NEW"],
 ["Evicted keys","Under memory pressure","> 0 trending = warn","NEW"],
],[2.0,2.4,1.9,0.6])

doc.add_page_break()
doc.add_heading('5. High-value attack signals (must-page)', level=1)
for b in [
 "100+ failed SSH logins per minute (brute force).",
 "Sudden outbound traffic spike from a server process (C2 / exfiltration).",
 "New process running as root / shell spawned in container.",
 "Unexpected new listening port or exposed DB/Redis port.",
 "Disk or inode > 90%; CPU > 90% sustained; swap > 50%.",
 "conntrack or file-descriptors > 85% (silent connection drops).",
 "New/modified cronjob or systemd unit (persistence).",
 "TLS certificate expiring in < 7 days, or already expired.",
 "Critical CVE in a deployed image or dependency.",
 "Changes to /etc/passwd, /etc/shadow, or sudoers.",
]: bullet(b)

doc.add_heading('6. Tooling map', level=1)
add_table(["Tool","Purpose","Where in repo"],[
 ["Your agent + TimescaleDB","Core metrics + security events","agent/, backend/, database/"],
 ["metrics-extended.js","Swap, inode, I/O, conntrack, FD, TCP, procs, drift","agent/src/collectors/"],
 ["cert-cve.js","TLS expiry, dep/OS CVEs","agent/src/collectors/"],
 ["fpm.js","PHP-FPM pool + top-CPU/mem worker & request","agent/src/collectors/"],
 ["Falco","Runtime/syscall + container threats","monitoring/security-tools/falco/"],
 ["Fail2Ban","Active brute-force blocking","monitoring/security-tools/fail2ban/"],
 ["auditd","OS-level audit events","monitoring/security-tools/auditd/"],
 ["Trivy","Image + filesystem CVE/secret scan","monitoring/security-tools/scan/"],
 ["gitleaks","Secret scanning in CI","monitoring/security-tools/ci/"],
 ["Lynis","Host hardening audit","run: lynis audit system"],
 ["OpenTelemetry","Distributed tracing","monitoring/app-instrumentation/"],
 ["Prometheus/Grafana/Tempo","Optional sidecar (exporters, blackbox, traces)","monitoring/app-instrumentation/"],
],[1.8,3.2,2.3])
doc.add_paragraph("* Items marked Gap are recommended next steps not yet implemented: GeoIP/impossible-travel detection on auth events, and a network IDS (Suricata or Zeek). Also worth adding operationally: verified restore-tested backups and an incident-response runbook with on-call rotation.")

# Footer page numbers
sec = doc.sections[0]
ftr = sec.footer.paragraphs[0]
ftr.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = ftr.add_run("Monitoring & Security Reference  -  Page ")
run.font.size = Pt(8); run.font.color.rgb = GREY
fld = OxmlElement('w:fldSimple'); fld.set(qn('w:instr'), 'PAGE')
ftr._p.append(fld)

doc.save(sys.argv[1])
print("written", sys.argv[1])
